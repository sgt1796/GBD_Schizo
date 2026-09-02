from __future__ import annotations

import argparse
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "D9EAF7"
TABLE_HEADER = "F2F4F7"
LIGHT_ORANGE = "FCE4D6"
GRAY = "666666"
GREEN = "E2F0D9"
TABLE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _table_column_widths(frame: pd.DataFrame, columns: list[tuple[str, str]], widths=None) -> list[int]:
    if widths:
        raw = [float(value) for value in widths]
    else:
        raw = []
        for key, label in columns:
            values = [str(label)] + [str(value) for value in frame[key].head(80).fillna("")]
            length = max(len(value) for value in values)
            raw.append(0.75 + min(length, 42) / 30)
    minimum = 520
    available = TABLE_WIDTH_DXA - minimum * len(raw)
    total = sum(raw) or 1
    result = [minimum + int(available * value / total) for value in raw]
    result[-1] += TABLE_WIDTH_DXA - sum(result)
    return result


def set_table_geometry(table, widths_dxa: list[int]):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag, attrs in (
        ("tblW", {"w": str(TABLE_WIDTH_DXA), "type": "dxa"}),
        ("tblInd", {"w": "120", "type": "dxa"}),
        ("tblLayout", {"type": "fixed"}),
    ):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), value)
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    repeat_table_header(row)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def set_image_alt(inline_shape, text: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", text)
    doc_pr.set("title", text[:80])


def setup_document(title: str, subtitle: str | None = None, provisional: bool = False, show_blinded: bool = True) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.header_distance = Inches(0.492); section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]; style.font.name = "Calibri"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.bold = True; style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"; style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); set_run_font(r, size=18, color=DARK_BLUE, bold=True)
    if subtitle:
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle); set_run_font(r2, size=11, color=GRAY); r2.italic = True
    if show_blinded:
        p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("Blinded manuscript"); set_run_font(r3, size=11, color=GRAY, bold=True)
    if provisional:
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.left_indent=Inches(.05); q.paragraph_format.right_indent=Inches(.05)
        q.paragraph_format.space_before=Pt(4); q.paragraph_format.space_after=Pt(10)
        q_pr=q._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),LIGHT_ORANGE); q_pr.append(shd)
        run = q.add_run(
            "PROVISIONAL ANALYTICAL DRAFT - NOT FOR SUBMISSION\n"
            "This build lacks the matching fine-age GBD 2023 burden and official population exports with verified provenance. "
            "It is suitable for analytical review only. Trend estimates are descriptive analyses of annual GBD posterior means."
        )
        run.bold = True; run.font.color.rgb = RGBColor(156, 87, 0)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return doc


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_dataframe_table(doc: Document, frame: pd.DataFrame, columns: list[tuple[str, str]], font_size=8.2, widths=None):
    frame = frame.copy()
    width_values = _table_column_widths(frame, columns, widths)
    table = doc.add_table(rows=1, cols=len(columns)); table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]; repeat_table_header(hdr)
    for i, (_, label) in enumerate(columns):
        cell = hdr.cells[i]; cell.text = label; set_cell_shading(cell, TABLE_HEADER); set_cell_margins(cell, start=120, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=font_size, color=DARK_BLUE, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, (key, _) in enumerate(columns):
            value = row.get(key, "")
            cells[i].text = "" if pd.isna(value) else str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cells[i], start=120, end=120)
            for run in cells[i].paragraphs[0].runs: set_run_font(run, size=font_size)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, width_values)
    doc.add_paragraph()
    return table


def add_figure(doc: Document, path: Path, caption: str, alt: str, width=6.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(width)); set_image_alt(shape, alt)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lead, rest = caption.split(". ", 1)
    cap.add_run(lead + ". ").bold = True; cap.add_run(rest)
    cap.paragraph_format.space_after = Pt(10)


def fmt(value, digits=1):
    if pd.isna(value): return "NA"
    return f"{float(value):,.{digits}f}"


def endpoint_lookup(endpoints: pd.DataFrame, loc: str, sex: str, outcome: str, metric: str):
    return endpoints[(endpoints.location_name==loc)&(endpoints.sex_name==sex)&(endpoints.measure_name==outcome)&(endpoints.metric_name==metric)].iloc[0]


def compact_endpoint_table(endpoints: pd.DataFrame, segmented: pd.DataFrame) -> pd.DataFrame:
    rows=[]
    for loc in ("China","United States of America"):
        for sex in ("Female","Male"):
            for outcome in ("Incidence","Prevalence","DALYs"):
                count=endpoint_lookup(endpoints,loc,sex,outcome,"All-age count")
                rate=endpoint_lookup(endpoints,loc,sex,outcome,"Age-standardized rate per 100,000")
                seg=segmented[(segmented.location_name==loc)&(segmented.sex_name==sex)&(segmented.measure_name==outcome)].iloc[0]
                rows.append({"Location":loc.replace("United States of America","United States"),"Sex":sex,"Outcome":outcome,
                             "2023 quantity (95% UI)":f"{fmt(count.value_2023,0)} ({fmt(count.lower_2023,0)}-{fmt(count.upper_2023,0)})",
                             "Quantity change, %":fmt(count.percent_change_point_estimate,1),
                             "ASR 1990":fmt(rate.value_1990,2),"ASR 2023 (95% UI)":f"{fmt(rate.value_2023,2)} ({fmt(rate.lower_2023,2)}-{fmt(rate.upper_2023,2)})",
                             "Descriptive AAPC, %/year":fmt(seg.aapc,3)})
    return pd.DataFrame(rows)


def build_manuscript(analysis: Path, out_dir: Path, meta: dict, tables: dict[str,pd.DataFrame]):
    provisional=not bool(meta.get("submission_ready", False))
    doc=setup_document("Contrasting sex-specific trends and demographic contributions to schizophrenia burden in China and the United States, 1990–2023",
                       "A comparative analysis of Global Burden of Disease 2023 estimates",provisional)
    endpoints=tables["endpoint_summary"]; segmented=tables["segmented_summary"]
    pairs=tables["trajectory_contrasts"]
    decomp=tables["decomposition"]

    add_heading(doc,"Abstract")
    china_f=endpoint_lookup(endpoints,"China","Female","DALYs","Age-standardized rate per 100,000")
    china_m=endpoint_lookup(endpoints,"China","Male","DALYs","Age-standardized rate per 100,000")
    us_f=endpoint_lookup(endpoints,"United States of America","Female","DALYs","Age-standardized rate per 100,000")
    us_m=endpoint_lookup(endpoints,"United States of America","Male","DALYs","Age-standardized rate per 100,000")
    add_paragraph(doc,"Background: Schizophrenia produces substantial lifelong disability, but comparisons of national trends can be obscured by population growth, population ageing, and changes in age-specific rates. We compared sex-specific schizophrenia burden in China and the United States, two populous settings with contrasting demographic trajectories.")
    add_paragraph(doc,"Methods: We analyzed GBD 2023 incidence, prevalence, and disability-adjusted life-year (DALY) estimates for 1990-2023. We summarized all-age quantities and age-standardized rates (ASRs), described nonlinear trends with Bayesian-information-criterion-selected segmented log-linear regression, and decomposed all-age changes into population-size, age-structure, and age-specific-rate components using Shapley replacement. Incidence age-period-cohort (APC) estimable functions were secondary. All derived results are point estimates because posterior draws were unavailable.")
    age_flags=int(tables["decomposition_age_bin_sensitivity"].material_age_bin_sensitivity.sum())
    add_paragraph(doc,f"Results: From 1990 to 2023, estimated incident cases, prevalent cases, and DALYs increased in every country-sex stratum. DALY ASRs in China changed from {fmt(china_f.value_1990,1)} to {fmt(china_f.value_2023,1)} per 100,000 among females and {fmt(china_m.value_1990,1)} to {fmt(china_m.value_2023,1)} among males. Corresponding U.S. rates changed from {fmt(us_f.value_1990,1)} to {fmt(us_f.value_2023,1)} among females and {fmt(us_m.value_1990,1)} to {fmt(us_m.value_2023,1)} among males. Descriptive segmented trends were near stable in China and declined in the United States. Population-size change was the principal positive contributor in most strata, but collapsing age bins materially changed at least one decomposition component in {age_flags} of 12 panels.")
    add_paragraph(doc,"Conclusions: Increasing incident and prevalent case counts and DALYs did not imply worsening age-standardized burden. The estimates identify contrasting demographic and rate patterns that may inform sex- and age-responsive capacity planning; they do not establish effects of health systems or policies.")
    add_paragraph(doc,"Keywords: schizophrenia; Global Burden of Disease; China; United States; segmented regression; decomposition; demographic change; sex differences")

    add_heading(doc,"Background")
    add_paragraph(doc,"Schizophrenia is a severe mental disorder associated with persistent functional impairment, premature mortality, family burden, and substantial health and social care needs [1-3]. In these GBD extracts, schizophrenia DALYs equal years lived with disability (YLDs) because GBD does not assign schizophrenia a direct fatal component; this accounting convention does not imply an absence of excess mortality among people with schizophrenia [1,3].")
    add_paragraph(doc,"Earlier global GBD analyses found that incident and prevalent case counts and DALYs increased substantially while age-standardized schizophrenia rates changed much less [2,3]. Counts and standardized rates therefore answer different questions: absolute quantities reflect population size and age structure as well as age-specific rates, whereas ASRs remove differences in population age composition.")
    add_paragraph(doc,"A 2026 GBD 2023 study of China already reported the same 1990-2023 endpoints, sex strata, and decomposition into population growth, ageing, and epidemiological change [4]. The Chinese 1990 and 2023 endpoint point estimates used here overlap exactly with that publication because both analyses use the same GBD release. GBD 2021 studies have also combined schizophrenia trends with sex and age patterns, country comparisons, decomposition, joinpoint models, inequalities, and projections [5-7]. Country-level trend and decomposition analysis is therefore a crowded literature, and neither method is claimed as novel here. The incremental contribution is a release-matched, sex-specific China-U.S. comparison through 2023 in one reproducible framework, with descriptive paired contrasts and no causal interpretation of ecological differences.")
    add_paragraph(doc,"We therefore examined incidence, prevalence, and DALY burden in China and the United States from 1990 through 2023. Our objectives were to quantify endpoints and nonlinear trends, compare country and sex trajectories descriptively, decompose all-age changes, and assess agreement between primary trends and secondary age-period-cohort estimable summaries.")

    add_heading(doc,"Methods")
    add_heading(doc,"Study design and data source",2)
    add_paragraph(doc,"We conducted a comparative secondary analysis of GBD 2023 modeled health estimates [8]. The analytic population comprised females and males in China and the United States from 1990 through 2023. The burden extract contained annual posterior means and 95% uncertainty intervals (UIs) for incidence, prevalence, YLDs, and DALYs by age, sex, country, year, and metric. The source files contained aggregated, non-identifiable population estimates; institutional review board review and informed consent were therefore not applicable.")
    add_heading(doc,"Outcomes and analytical exclusions",2)
    add_paragraph(doc,"Primary outcomes were incidence, prevalence, and DALYs. Trend analyses used all-age numbers and ASRs per 100,000. The production decomposition contract requires matching Number and Rate estimates plus population for 20 mutually exclusive age groups from 0-4 through 95+ years. The available provisional extract instead supports 13 groups from 0-14 through 70+, so every result in this build remains non-submission. We excluded the GBD Percent metric because its denominator and age basis differed by outcome. Probability-of-death and available risk-factor extracts were excluded because they did not represent schizophrenia-specific causal attribution. YLD and DALY estimates were audited for numerical identity and YLDs were not duplicated in results.")
    add_heading(doc,"Data quality and population denominators",2)
    source_text="official GBD 2023 population estimates" if meta["population_status"]=="official_GBD_2023" else "a provisional population proxy reconstructed from matched GBD count-rate pairs"
    add_paragraph(doc,f"We audited dimensional uniqueness, completeness across 34 years, missingness, positivity, uncertainty-bound ordering, and age-bin consistency. Decomposition currently uses {source_text}. A submission build requires both a matching fine-age burden export and an authenticated population export from GBD 2023, with retrieval records, export identifiers, query dimensions, preserved raw files, and verified SHA-256 hashes. We checked age-specific reconstruction by multiplying population by rate and dividing by 100,000.")
    add_heading(doc,"Trend and comparison analyses",2)
    add_paragraph(doc,"We modeled log ASRs with an open Python implementation of piecewise linear regression, using the general joinpoint parameterization [9]. Candidate models allowed zero to two change points and required at least four annual observations per segment. The number and locations of change points were selected by the Bayesian information criterion. Segment annual percentage changes and overall average annual percentage changes (AAPCs) summarize fitted changes in the annual GBD posterior means; they are descriptive and are not presented with model p values or confidence intervals. Sensitivities varied the maximum breakpoint count, minimum segment length, calendar window, rate versus log-rate scale, and weights based on log-UI widths. The native UIs are not sampling standard errors, so weighted results also remained descriptive. A registered NCI Joinpoint 6.1.0 analysis may be used as optional validation but is not required for the primary open workflow [10].")
    add_paragraph(doc,"We summarized China-U.S. contrasts within outcome and sex and female-male contrasts within outcome and country as differences in descriptive AAPCs. Lag-1 residual autocorrelation was used as a diagnostic because annual GBD estimates may be serially correlated. We did not conduct parallelism F tests or multiplicity-adjusted significance tests. Native GBD UIs were reported only for original GBD estimates. Derived ratios, differences, changes, AAPCs, and decomposition components remained point estimates because posterior draws and cross-estimate correlations were unavailable.")
    add_heading(doc,"Demographic decomposition",2)
    add_paragraph(doc,"For each outcome, country, and sex, we decomposed the change in reconstructed all-age quantities into population-size change, age-structure change, and age-specific-rate change [11]. We averaged marginal contributions over all six possible factor-replacement orders, which is equivalent to a Shapley decomposition. Components were required to sum to total reconstructed change within numerical tolerance, and reconstructed endpoint quantities were checked against the reported all-age numbers. Primary estimates used 1990-2023; 2000-2023, 2010-2023, annual-chain, and five-year-chain analyses assessed endpoint and path sensitivity. We also compared the finest available age partition with a four-group collapse and flagged sign, magnitude-rank, or component-size changes.")
    add_heading(doc,"Secondary age-period-cohort analysis",2)
    add_paragraph(doc,"Secondary incidence analysis used equal five-year age groups from 15-19 through 65-69 years and six equal periods from 1994-1998 through 2019-2023. The production contract will add ages 10-14 when available. We estimated net drift, age-specific local drift, longitudinal age relative risks, and nonlinear period and cohort relative risks. The design included an intercept, two identifiable linear trends, and nonlinear age, period, and cohort bases that were orthogonal to intercept and linear trend; it did not fit three unrestricted linear effects [12-15]. Reference relative risks were set to one. We did not interpret period or cohort patterns causally. The grouped sensitivity used the six complete periods from 1990-1994 through 2015-2019, thereby excluding 2020-2023 without shortening the six-period design.")
    add_heading(doc,"Reporting and reproducibility",2)
    add_paragraph(doc,"The analysis was conducted in Python with deterministic settings. Every table and figure is generated from saved machine-readable inputs. Reporting follows the Guidelines for Accurate and Transparent Health Estimates Reporting (GATHER) [16]. The working package contains the complete code, provenance table, data dictionary, and readiness metadata; these materials will be deposited at a persistent repository before submission. OpenAI Codex, a large language model tool, was used during software engineering and manuscript drafting. All generated code, analyses, claims, and references were reviewed and remain the authors' responsibility; the tool is not an author.")

    add_heading(doc,"Results")
    add_heading(doc,"Data completeness and outcome selection",2)
    audit=tables["data_audit"]; ident=tables["yld_daly_identity"].iloc[0]; recon=tables["population_reconstruction"]
    age_count=int(decomp.age_group_count.iloc[0])
    add_paragraph(doc,f"All 12 primary country-sex-outcome panels contained 34 annual all-age quantities and 34 annual ASRs. The provisional decomposition panels contained {age_count} age groups spanning 0-14 through 70+ years. No invalid UI ordering or nonpositive point estimates were identified. DALYs and YLDs were numerically identical across {int(ident.matched_cells):,} matched cells (maximum relative point-estimate difference {ident.max_relative_difference_val:.2e}); YLDs were therefore omitted as duplicate outcomes. The 99th percentile absolute population-rate reconstruction discrepancy was {recon.relative_error_pct.abs().quantile(.99):.3g}%, and decomposition endpoint reconstructions agreed with reported all-age numbers within floating-point tolerance.")
    add_heading(doc,"Burden levels and temporal changes",2)
    add_paragraph(doc,"All-age incidence, prevalence, and DALY counts increased between 1990 and 2023 in both countries and both sexes. In contrast, Chinese ASRs were nearly stable, whereas U.S. ASRs declined across the three outcomes. Males generally had higher rates than females. Thus, count growth and standardized-rate trends conveyed different dimensions of population burden.")
    table1=compact_endpoint_table(endpoints,segmented)
    add_paragraph(doc,"Table 1. Endpoint burden and descriptive trends",bold_lead="Table 1.")
    add_dataframe_table(doc,table1,[("Location","Location"),("Sex","Sex"),("Outcome","Outcome"),("2023 quantity (95% UI)","2023 quantity (95% UI)"),("Quantity change, %","Quantity change, %"),("ASR 1990","ASR 1990"),("ASR 2023 (95% UI)","ASR 2023 (95% UI)"),("Descriptive AAPC, %/year","Descriptive AAPC, %/year")],font_size=7.2)
    add_paragraph(doc,"Quantity denotes incident cases, prevalent cases, or DALYs; DALYs are healthy life-years lost, not people. The 95% UIs apply to native GBD endpoint estimates. Percentage changes and AAPCs are point summaries without propagated posterior uncertainty.")
    add_figure(doc,analysis/"figures"/"main"/"figure_1_asr_trends.png","Figure 1. Sex-specific schizophrenia age-standardized rates in China and the United States, 1990-2023. Shading shows native 95% GBD uncertainty intervals.","Three panels showing incidence, prevalence, and DALY age-standardized rate trends by country and sex.")
    add_heading(doc,"Segmented trends and descriptive comparisons",2)
    add_paragraph(doc,"BIC-selected segmented curves summarized nonlinear features of the annual posterior means. Residual diagnostics indicated that an independent-error approximation was not uniformly plausible, reinforcing descriptive interpretation. No change point, slope, or between-group contrast was interpreted as a significance test.")
    country_pairs=pairs[pairs.comparison_family=="country"].copy()
    country_pairs["China change"]=country_pairs.group_a_annualized_endpoint_change_pct.map(lambda x:fmt(x,3))
    country_pairs["US change"]=country_pairs.group_b_annualized_endpoint_change_pct.map(lambda x:fmt(x,3))
    country_pairs["Difference"]=country_pairs.annualized_endpoint_change_difference_b_minus_a_pct_points.map(lambda x:fmt(x,3))
    add_paragraph(doc,"Table 2. China-United States annualized endpoint contrasts",bold_lead="Table 2.")
    add_dataframe_table(doc,country_pairs,[("stratum","Sex"),("measure_name","Outcome"),("China change","China, %/year"),("US change","United States, %/year"),("Difference","US-China, percentage points")],font_size=8)
    add_paragraph(doc,"Annualized endpoint changes are descriptive point estimates calculated from the 1990 and 2023 ASRs. The difference is United States minus China; no parallelism test or interval is implied. Full country and sex contrasts are reported in the supplement.")
    add_figure(doc,analysis/"figures"/"main"/"figure_2_segmented_trends.png","Figure 2. Observed and fitted descriptive segmented age-standardized rate trajectories. Curves were selected by BIC in the open Python workflow and are not NCI Joinpoint output.","Six panels showing observed and fitted sex-specific segmented rate trends.")
    add_heading(doc,"Age-specific patterns",2)
    add_paragraph(doc,"The age profiles differed by outcome and sex, but male rates were generally higher. Incidence peaked in early adulthood, whereas prevalence and DALY rates remained substantial across middle adulthood. The coarse 70+ terminal category limited interpretation of late-life heterogeneity.")
    add_figure(doc,analysis/"figures"/"main"/"figure_3_age_patterns.png","Figure 3. Age-specific incidence, prevalence, and DALY rates in 2023 by country and sex. Shading shows native 95% GBD uncertainty intervals.","Six panels showing age-specific schizophrenia rates in 2023 by country and sex.")
    add_heading(doc,"Drivers of changing all-age quantities",2)
    primary=decomp[(decomp.start_year==1990)&(decomp.end_year==2023)].copy()
    primary["Location"]=primary.location_name.str.replace("United States of America","United States",regex=False)
    for c in ("population_size_change","age_structure_change","age_specific_rate_change","total_change"): primary[c]=primary[c].map(lambda x:fmt(x,0))
    add_paragraph(doc,"Table 3. Shapley decomposition of all-age changes, 1990-2023",bold_lead="Table 3.")
    add_dataframe_table(doc,primary,[("Location","Location"),("sex_name","Sex"),("measure_name","Outcome"),("population_size_change","Population size"),("age_structure_change","Age structure"),("age_specific_rate_change","Age-specific rate"),("total_change","Total change")],font_size=8)
    add_paragraph(doc,"Population-size change contributed positively in every panel. In China, age-structure change reduced incident cases but increased prevalent cases and DALYs; age-specific-rate change was small relative to demographic components. In the United States, age-structure and rate changes countervailed population-size change across all three outcomes and both sexes. Negative components represent countervailing forces; component percentages can therefore exceed 100% and were not treated as compositional shares. The age-bin sensitivity flagged material variation in selected incidence panels, so component magnitudes from this provisional broad-bin build should not be interpreted as final.")
    add_figure(doc,analysis/"figures"/"main"/"figure_4_decomposition.png","Figure 4. Shapley decomposition of changes in reconstructed all-age schizophrenia quantities, 1990-2023. Components are deterministic attributions based on posterior mean rates.","Six panels decomposing all-age changes into population-size, age-structure, and age-specific-rate components.",width=5.55)
    add_heading(doc,"Secondary age-period-cohort summaries",2)
    if "apc_primary_direction_agreement" in tables:
        direction=tables["apc_primary_direction_agreement"]
        agreement_count=int(direction.apc_vs_segmented_direction_agreement.fillna(False).sum())
        agreement_total=int(direction.apc_vs_segmented_direction_agreement.notna().sum())
    else:
        direction=tables["apc_summary"][["location_name","sex_name","net_drift"]].merge(
            segmented[segmented.measure_name=="Incidence"][["location_name","sex_name","aapc"]],
            on=["location_name","sex_name"],validate="one_to_one")
        agreement_count=int(((direction.net_drift>0)==(direction.aapc>0)).sum())
        agreement_total=len(direction)
    add_paragraph(doc,f"The signs of incidence APC net drift and primary segmented AAPC agreed in {agreement_count} of {agreement_total} country-sex strata, indicating partial rather than uniform directional corroboration. This is not expected to be exact numerical agreement because ASR segmented trends and selected-age APC net drift use different estimands, age coverage, standardization, weighting, model form, calendar aggregation, and identification constraints. Longitudinal age and nonlinear period and cohort relative risks are presented in the supplement and normalized to one at their reference categories. Interpretation remained descriptive because the exact linear dependency among age, period, and cohort prevents unique causal separation.")

    add_heading(doc,"Discussion")
    add_paragraph(doc,"This comparative study found increases in estimated incident cases, people living with schizophrenia, and DALYs, alongside contrasting standardized-rate trajectories. Chinese ASRs were broadly stable, whereas U.S. ASRs declined, particularly among females. The combination of increasing absolute quantities and flat or falling ASRs is not contradictory: population growth, age composition, and age-specific rates contribute separately.")
    add_paragraph(doc,"The China endpoint point estimates are identical to those reported by Song and colleagues because both studies use the same GBD 2023 source [4]; this is expected data overlap, not independent validation. The component pattern also closely reproduces that study: population growth was the largest positive component, ageing reduced incidence but increased prevalence and DALYs, and rate change was comparatively small. The added comparison shows a different U.S. pattern, in which population growth was partly offset by age-composition and rate changes for each outcome and sex. These deterministic decompositions allocate endpoint differences under the specified identity; they do not identify causes of the demographic or rate changes.")
    add_paragraph(doc,"The distinction between measures has practical relevance. Incident and prevalent case counts can help approximate potential service-volume need, whereas DALYs quantify healthy life lost and are not a count of people. None of these quantities measures realized service use or demand, which also depends on detection, access, coverage, severity, and treatment pathways. ASRs facilitate temporal and cross-country rate comparisons after standardizing age composition. Planning based on ASRs alone may miss growth in potential service volume, while interpreting count growth as worsening individual risk would also be misleading.")
    add_paragraph(doc,"Sex-specific estimates showed a persistent male excess for several outcomes and an incidence peak in younger age groups, patterns broadly consistent with earlier GBD schizophrenia analyses [2,3]. However, modeled estimates cannot determine whether country or sex contrasts reflect underlying incidence, diagnostic recognition, data availability, care access, remission, excess mortality, or model assumptions. Health-system and policy differences remain contextual hypotheses rather than tested mechanisms.")
    add_heading(doc,"Strengths and limitations",2)
    add_paragraph(doc,"Strengths include a focused comparative question, consistent GBD 2023 outcome definitions, explicit separation of native UIs from derived point summaries, deterministic and reproducible trend selection, exact Shapley decomposition, a dedicated APC module with synthetic recovery tests, and restriction of APC interpretation to identifiable estimable functions. The pipeline records exclusions and prevents incomplete age or provenance inputs from being silently treated as submission-ready.")
    add_paragraph(doc,"Several limitations are important. First, GBD values are modeled estimates rather than direct observations and may share smoothing assumptions across years, sexes, outcomes, and countries. Annual posterior means therefore cannot be treated as independent observations; segmented curves, change points, AAPCs, and pairwise contrasts are descriptive. Second, posterior draws were unavailable, so uncertainty could not be propagated to changes, ratios, decomposition components, or trajectory contrasts. Third, the provisional broad 0-14 and 70+ bins obscure heterogeneity within childhood and late life, and the explicit collapse sensitivity materially changed selected decomposition results. Fourth, APC estimates are restricted to incidence at ages 15-69 in this extract, are sensitive to grouping, and cannot identify independent causal age, period, and cohort effects. Comparisons with ASR trends are qualitative because the estimands differ. Fifth, ecological country contrasts cannot support causal attribution to policy or health systems. Finally, the provisional build lacks both matching fine-age burden data and an authenticated GBD 2023 population export with complete provenance. Registered NCI Joinpoint output and posterior draws would provide useful validation and uncertainty propagation, respectively, but are not mandatory for the primary open analysis.")
    add_heading(doc,"Conclusions")
    add_paragraph(doc,"Between 1990 and 2023, estimated incident and prevalent case counts and DALYs increased in China and the United States while standardized-rate trajectories differed. Population growth contributed positively throughout, but age-composition and rate components contrasted between countries. These modeled patterns support considering potential service volume, health loss, standardized rates, age structure, and sex jointly in planning; they do not establish causal effects of policy or health systems.")

    add_heading(doc,"List of abbreviations")
    add_paragraph(doc,"AAPC: average annual percentage change; APC: age-period-cohort; ASR: age-standardized rate; BIC: Bayesian information criterion; DALY: disability-adjusted life-year; GATHER: Guidelines for Accurate and Transparent Health Estimates Reporting; GBD: Global Burden of Disease; IHME: Institute for Health Metrics and Evaluation; NCI: National Cancer Institute; UI: uncertainty interval; YLD: year lived with disability.")

    add_heading(doc,"Declarations")
    add_heading(doc,"Ethics approval and consent to participate",2); add_paragraph(doc,"Not applicable. This study used aggregated, non-identifiable modeled estimates available through the Institute for Health Metrics and Evaluation.")
    add_heading(doc,"Consent for publication",2); add_paragraph(doc,"Not applicable.")
    add_heading(doc,"Availability of data and materials",2); add_paragraph(doc,"GBD estimates are available through the IHME GBD Results Tool subject to IHME terms [8]. Analytic code, derived tables, provenance metadata, and exact model settings are prepared for archiving, but a persistent repository link must be added before submission. The repository will not redistribute source data beyond applicable IHME terms.")
    add_heading(doc,"Competing interests",2); add_paragraph(doc,"The authors declare no competing interests.")
    add_heading(doc,"Funding",2); add_paragraph(doc,"Funding for this secondary analysis and any funder roles must be confirmed by the authors before submission.")
    add_heading(doc,"Authors' contributions",2); add_paragraph(doc,"Contributor roles will be reported using the CRediT taxonomy in the unblinded submission file.")
    add_heading(doc,"Acknowledgements",2); add_paragraph(doc,"We acknowledge the Institute for Health Metrics and Evaluation and the Global Burden of Disease collaborators for producing and making the GBD 2023 estimates available. This acknowledgment does not imply their endorsement of the present secondary analysis.")

    add_heading(doc,"References")
    refs=[
        "1. GBD 2019 Mental Disorders Collaborators. Global, regional, and national burden of 12 mental disorders in 204 countries and territories, 1990-2019: a systematic analysis for the Global Burden of Disease Study 2019. Lancet Psychiatry. 2022;9(2):137-150. doi:10.1016/S2215-0366(21)00395-3.",
        "2. Charlson FJ, Ferrari AJ, Santomauro DF, et al. Global epidemiology and burden of schizophrenia: findings from the Global Burden of Disease Study 2016. Schizophr Bull. 2018;44(6):1195-1203. doi:10.1093/schbul/sby058.",
        "3. Solmi M, Seitidis G, Mavridis D, et al. Incidence, prevalence, and global burden of schizophrenia - data, with critical appraisal, from the Global Burden of Disease (GBD) 2019. Mol Psychiatry. 2023;28(12):5319-5327. doi:10.1038/s41380-023-02138-4.",
        "4. Song H, Fang L, Li H. Long-term trends in the burden of schizophrenia in China, 1990-2052: an analysis of the Global Burden of Disease Study 2023. Psychol Res Behav Manag. 2026;19:574531. doi:10.2147/PRBM.S574531.",
        "5. Huo J, Li R, Ren X, et al. Trends in incidence, prevalence, and disability-adjusted life years of schizophrenia in China from 1990 to 2021, with projections for 2022-2050. Front Psychiatry. 2025;16:1651350. doi:10.3389/fpsyt.2025.1651350.",
        "6. Luo W, Gao J, Guo Z, et al. Trends and cross-country inequalities in schizophrenia from 1990 to 2021, with prediction to 2035: a systematic analysis of the Global Burden of Disease Study 2021. BMC Psychiatry. 2025;25:928. doi:10.1186/s12888-025-07273-6.",
        "7. Zhou W, He J, Wu L, et al. Global, regional, and national burden of schizophrenia: epidemiological trends, decomposition, joinpoint analysis, and projections to 2036 based on GBD 2021. Front Psychiatry. 2026;17:1702808. doi:10.3389/fpsyt.2026.1702808.",
        "8. Global Burden of Disease Collaborative Network. Global Burden of Disease Study 2023 (GBD 2023) Results. Seattle, United States: Institute for Health Metrics and Evaluation; 2024. https://vizhub.healthdata.org/gbd-results/. Accessed 1 September 2026.",
        "9. Kim HJ, Fay MP, Feuer EJ, Midthune DN. Permutation tests for joinpoint regression with applications to cancer rates. Stat Med. 2000;19(3):335-351; correction 2001;20(4):655.",
        "10. National Cancer Institute. Joinpoint Regression Program, version 6.1.0. Surveillance Research Program. https://surveillance.cancer.gov/joinpoint/. Accessed 1 September 2026.",
        "11. Das Gupta P. Standardization and decomposition of rates: a user's manual. Current Population Reports, Series P23-186. Washington, DC: US Bureau of the Census; 1993.",
        "12. Holford TR. The estimation of age, period and cohort effects for vital rates. Biometrics. 1983;39(2):311-324.",
        "13. Clayton D, Schifflers E. Models for temporal variation in cancer rates. II: age-period-cohort models. Stat Med. 1987;6(4):469-481.",
        "14. Luo L. Assessing validity and application scope of the intrinsic estimator approach to the age-period-cohort problem. Demography. 2013;50(6):1945-1967.",
        "15. Rutherford MJ, Lambert PC, Thompson JR. Age-period-cohort modeling. Stata J. 2010;10(4):606-627.",
        "16. Stevens GA, Alkema L, Black RE, et al. Guidelines for Accurate and Transparent Health Estimates Reporting: the GATHER statement. Lancet. 2016;388(10062):e19-e23. doi:10.1016/S0140-6736(16)30388-9.",
    ]
    for ref in refs:
        p_ref=doc.add_paragraph(ref,style=None)
        p_ref.paragraph_format.first_line_indent=Inches(-.2); p_ref.paragraph_format.left_indent=Inches(.2)
        p_ref.paragraph_format.space_after=Pt(3)
        for run in p_ref.runs: run.font.size=Pt(9.3)

    out=out_dir/"manuscript_BMC_Public_Health.docx"; doc.save(out); return out


def build_supplement(analysis: Path, out_dir: Path, meta: dict, tables: dict[str,pd.DataFrame]):
    provisional=not bool(meta.get("submission_ready", False))
    doc=setup_document("Supplementary material","China-US schizophrenia burden study, 1990-2023",provisional)
    add_heading(doc,"Supplementary methods and audit trail")
    add_paragraph(doc,"This supplement contains the full audit, sensitivity analyses, secondary APC outputs, and readiness conditions. Native GBD uncertainty intervals are never relabeled as confidence intervals for derived quantities.")
    add_heading(doc,"Table S1. Data provenance",2)
    provenance=tables["provenance"].copy()
    provenance["file"]=provenance["file"].map(lambda x:Path(str(x)).name if ("/" in str(x) or "\\" in str(x)) else str(x))
    provenance["status"]=provenance["status"].replace({"derived_proxy_NOT_OFFICIAL":"PROVISIONAL proxy"})
    provenance_columns=[("source_role","Source role"),("gbd_release","GBD release"),("file","File"),("retrieval_date","Retrieval date"),("export_id","Export ID"),("status","Status")]
    add_dataframe_table(doc,provenance,provenance_columns,font_size=7.8)
    add_paragraph(doc,"Complete SHA-256 hashes, metadata-file hashes, and query dimensions are retained in provenance.csv and publication_tables.xlsx.")
    doc.add_page_break()
    add_heading(doc,"Table S2. Completeness and validity audit",2)
    add_dataframe_table(doc,tables["data_audit"],[(c,c.replace("_"," ").title()) for c in tables["data_audit"].columns],font_size=7.2)
    add_heading(doc,"Table S3. YLD-DALY identity audit",2)
    ident=tables["yld_daly_identity"].copy()
    for c in ident.columns:
        if c.startswith("max_"): ident[c]=ident[c].map(lambda v:f"{float(v):.3e}")
    add_dataframe_table(doc,ident,[(c,c.replace("_"," ").title()) for c in ident.columns],font_size=7.5)
    doc.add_page_break()
    add_heading(doc,"Table S4. Full descriptive trajectory contrasts",2)
    contrasts=tables["trajectory_contrasts"].copy()
    contrasts["group_a"]=contrasts.group_a.str.replace("United States of America","United States",regex=False)
    contrasts["group_b"]=contrasts.group_b.str.replace("United States of America","United States",regex=False)
    contrasts["stratum"]=contrasts.stratum.str.replace("United States of America","United States",regex=False)
    contrasts["Comparison"]=contrasts.group_a+" vs "+contrasts.group_b
    contrasts["Group A change"]=contrasts.group_a_annualized_endpoint_change_pct.map(lambda x:fmt(x,3))
    contrasts["Group B change"]=contrasts.group_b_annualized_endpoint_change_pct.map(lambda x:fmt(x,3))
    contrasts["Difference"]=contrasts.annualized_endpoint_change_difference_b_minus_a_pct_points.map(lambda x:fmt(x,3))
    contrasts["Same direction"]=contrasts.same_annualized_endpoint_change_direction.map(lambda x:"Yes" if bool(x) else "No")
    contrast_columns=[("comparison_family","Family"),("stratum","Stratum"),("measure_name","Outcome"),("Comparison","Comparison"),("Group A change","Group A, %/year"),("Group B change","Group B, %/year"),("Difference","B-A, percentage points"),("Same direction","Same direction")]
    add_dataframe_table(doc,contrasts,contrast_columns,font_size=7.3)
    add_paragraph(doc,"All entries are descriptive annualized endpoint contrasts; no hypothesis test or confidence interval is reported.")
    add_heading(doc,"Table S5. Weighted trend sensitivity",2)
    x=tables["ui_weighted_sensitivity"].copy()
    for c in ("primary_aapc","ui_weighted_fixed_knot_aapc","difference"): x[c]=x[c].map(lambda v:fmt(v,4))
    add_dataframe_table(doc,x,[(c,c.replace("_"," ").title()) for c in x.columns],font_size=7.5)
    add_heading(doc,"Table S6. Segmented-model specification sensitivity",2)
    spec=tables["segmented_specification_sensitivity"].copy()
    spec=spec[(spec.specification=="primary") | (~spec.direction_stable_vs_primary)].copy()
    spec["location_name"]=spec.location_name.str.replace("United States of America","United States",regex=False)
    spec["Direction stable"]=spec.direction_stable_vs_primary.map(lambda value:"Yes" if bool(value) else "No")
    spec["Change"]=spec.fitted_annualized_endpoint_change_pct.map(lambda value:fmt(value,3))
    spec_columns=[("location_name","Location"),("sex_name","Sex"),("measure_name","Outcome"),("specification","Specification"),("model_scale","Scale"),("Change","Annualized change, %"),("trajectory_direction","Direction"),("Direction stable","Stable vs primary")]
    add_dataframe_table(doc,spec,spec_columns,font_size=7.0)
    add_paragraph(doc,"The table shows each primary fit and every direction-unstable sensitivity. The complete specification grid is supplied in the machine-readable tables. Specifications varied breakpoint count, minimum segment length, calendar window, and model scale. Stability refers to the sign of the fitted endpoint trajectory, not inferential equivalence.")
    add_heading(doc,"Table S7. Alternative-window decomposition",2)
    d=tables["decomposition"][tables["decomposition"].start_year.isin([2000,2010])].copy()
    d["location_name"]=d.location_name.str.replace("United States of America","United States",regex=False)
    for c in ("population_size_change","age_structure_change","age_specific_rate_change","total_change"):
        if c in d: d[c]=d[c].map(lambda v:fmt(v,1))
    d["closure_error"]=d.closure_error.map(lambda v:"<1e-8" if abs(float(v))<1e-8 else f"{float(v):.2e}")
    add_dataframe_table(doc,d,[(c,c.replace("_"," ").title()) for c in ["location_name","sex_name","measure_name","start_year","end_year","population_size_change","age_structure_change","age_specific_rate_change","total_change","closure_error"]],font_size=7.2)
    add_heading(doc,"Table S8. Age-bin decomposition sensitivity",2)
    bins=tables["decomposition_age_bin_sensitivity"].copy()
    bins["location_name"]=bins.location_name.str.replace("United States of America","United States",regex=False)
    rank_labels={"population_size_change":"Population size","age_structure_change":"Age structure","age_specific_rate_change":"Age-specific rate"}
    for column in ("finest_component_rank","collapsed_component_rank"):
        bins[column]=bins[column].map(lambda value:" > ".join(rank_labels.get(part,part) for part in str(value).split(" > ")))
    bins["Material"]=bins.material_age_bin_sensitivity.map(lambda value:"Yes" if bool(value) else "No")
    bins["Maximum shift"]=bins.maximum_component_shift_pct_of_total_change.map(lambda value:fmt(value,1))
    bin_columns=[("location_name","Location"),("sex_name","Sex"),("measure_name","Outcome"),("finest_age_group_count","Finest bins"),("collapsed_age_group_count","Collapsed bins"),("finest_component_rank","Finest rank"),("collapsed_component_rank","Collapsed rank"),("Maximum shift","Maximum shift, % of total"),("Material","Material")]
    add_dataframe_table(doc,bins,bin_columns,font_size=7.1)
    add_paragraph(doc,"A result was flagged when a component changed sign, component magnitude ranking changed, or a component shifted by at least 10% of total change. The comparison is an aggregation sensitivity, not an uncertainty interval.")
    add_heading(doc,"Secondary APC analysis",1)
    add_paragraph(doc,"The APC analysis is intentionally restricted to incidence and identifiable summaries. The model uses two linear trends and nonlinear age, period, and cohort curvature bases rather than three unrestricted linear effects. Longitudinal age, period, and cohort relative risks are normalized to one at their reference categories. They are not separately identified causal effects.")
    add_heading(doc,"Table S9. Descriptive incidence APC net drift",2)
    apc=tables["apc_summary"].copy()
    apc["Location"]=apc.location_name.str.replace("United States of America","United States",regex=False)
    apc["Period"]=apc.start_year.astype(int).astype(str)+"-"+apc.end_year.astype(int).astype(str)
    apc["Net drift"]=apc.net_drift.map(lambda v:fmt(v,4))
    add_dataframe_table(doc,apc,[("Location","Location"),("sex_name","Sex"),("Period","Period"),("Net drift","Net drift, %/year")],font_size=8)
    add_paragraph(doc,"Net drift is a descriptive point estimate; no model confidence interval is reported.")
    if "apc_sensitivity_summary_1990_2019" in tables:
        add_heading(doc,"Table S10. APC sensitivity, 1990-2019",2)
        apc_sensitivity=tables["apc_sensitivity_summary_1990_2019"].copy()
        apc_sensitivity["Location"]=apc_sensitivity.location_name.str.replace("United States of America","United States",regex=False)
        apc_sensitivity["Period"]=apc_sensitivity.start_year.astype(int).astype(str)+"-"+apc_sensitivity.end_year.astype(int).astype(str)
        apc_sensitivity["Net drift"]=apc_sensitivity.net_drift.map(lambda v:fmt(v,4))
        add_dataframe_table(doc,apc_sensitivity,[("Location","Location"),("sex_name","Sex"),("Period","Period"),("Net drift","Net drift, %/year")],font_size=8)
        add_paragraph(doc,"This sensitivity retains six equal five-year periods, from 1990-1994 through 2015-2019, and therefore excludes 2020-2023 without shortening the grouped design.")
    if "apc_primary_direction_agreement" in tables:
        add_heading(doc,"Table S11. Primary-APC direction comparison",2)
        agreement=tables["apc_primary_direction_agreement"].copy()
        agreement["Location"]=agreement.location_name.str.replace("United States of America","United States",regex=False)
        agreement["Agrees with segmented"]=agreement.apc_vs_segmented_direction_agreement.map(lambda x:"Yes" if bool(x) else "No")
        agreement_columns=[("Location","Location"),("sex_name","Sex"),("primary_segmented_direction","Segmented direction"),("primary_observed_direction","Observed endpoint direction"),("apc_net_drift_direction","APC direction"),("Agrees with segmented","APC-segmented agreement")]
        add_dataframe_table(doc,agreement,agreement_columns,font_size=7.6)
        add_paragraph(doc,"Agreement denotes only matching signs of descriptive point estimates; it is not a significance test and numerical agreement is not expected because the estimands differ.")
    add_heading(doc,"Table S12. Cross-method contradiction audit",2)
    contradiction=tables["cross_method_contradictions"].copy()
    contradiction["location_name"]=contradiction.location_name.str.replace("United States of America","United States",regex=False)
    contradiction_columns=[("location_name","Location"),("sex_name","Sex"),("measure_name","Outcome"),("segmented_direction_1990_2023","Segmented"),("asr_endpoint_direction_1994_2023","ASR endpoint"),("selected_age_crude_direction_1994_2023","Selected-age crude"),("apc_net_drift_direction_1994_2023","APC net drift"),("likely_explanatory_factors","Likely factors")]
    add_dataframe_table(doc,contradiction,contradiction_columns,font_size=7.0)
    add_paragraph(doc,"Each apparent mismatch was checked on an aligned 1994-2023 calendar window. Likely factors include estimand, age coverage, standardization, weighting, model form, and APC constraints; the tested implementation did not indicate a software failure.")
    add_figure(doc,analysis/"figures"/"supplement"/"figure_s1_counts.png","Figure S1. All-age schizophrenia quantities by country and sex, 1990-2023. DALYs represent healthy life-years lost rather than people.","Three panels showing all-age incident cases, prevalent cases, and DALYs.")
    add_figure(doc,analysis/"figures"/"supplement"/"figure_s2_apc_incidence.png","Figure S2. Secondary incidence APC estimable summaries. Relative-risk curves are nonlinear curvature contrasts, not independent causal age, period, or cohort effects.","Four panels showing incidence net drift and nonlinear age, period, and cohort curvature summaries.")
    add_heading(doc,"Readiness condition and optional validation")
    add_bullets(doc,[
        f"Population status: {meta['population_status']}",
        f"Fine-age burden validated: {meta.get('fine_age_burden_validated', False)}",
        f"Source metadata complete: {meta.get('source_metadata_complete', False)}",
        f"Trend status: {meta['trend_status']}",
        "The matching fine-age burden and official population exports, each with a validated provenance sidecar and preserved raw files, are required before submission.",
        "Registered NCI Joinpoint output is optional validation; the primary trend workflow is the open descriptive BIC implementation and must not be described as NCI Joinpoint.",
        "Posterior draws are desirable for uncertainty propagation but are currently unavailable; derived quantities must remain explicitly labeled as point estimates.",
    ])
    out=out_dir/"supplementary_material.docx"; doc.save(out); return out


def build_methods_appendix(out_dir: Path, meta: dict):
    doc=setup_document("Statistical methods appendix","China-US schizophrenia burden study",False,show_blinded=False)
    if not bool(meta.get("submission_ready", False)):
        note=doc.add_paragraph()
        run=note.add_run("Provisional methods appendix: matching fine-age GBD 2023 burden and official population exports with validated provenance are required before submission.")
        run.bold=True; run.font.color.rgb=RGBColor(156,87,0)
    add_heading(doc,"Analysis estimands")
    add_bullets(doc,["Incident and prevalent case counts approximate potential service-volume need; DALYs quantify health loss and are not people.","Age-standardized rates describe temporal and cross-country rate patterns after standardization.","Derived ratios, changes, trend contrasts, APC functions, and decomposition components are point estimands; no draw-based uncertainty is asserted.","Production decomposition components are deterministic functions of posterior mean rates and populations across 20 fine-age groups from 0-4 through 95+ years; the provisional build uses 13 broader groups."])
    add_heading(doc,"Segmented log-linear trend model")
    add_paragraph(doc,"For year t and rate r(t), log r(t) = beta0 + beta1(t-t0) + sum[j] delta_j max(0,t-tau_j) + error(t). Candidate models contain zero, one, or two knots and require four observations between boundaries. The model with minimum BIC is selected, counting each knot location in the parameter penalty. Segment APC is 100[exp(slope)-1]; AAPC is the time-weighted mean log slope transformed to a percentage.")
    add_paragraph(doc,"The implementation is deterministic, open, and descriptive. Residual-permutation p values and regression confidence intervals are not used because serial dependence and the joint GBD posterior covariance are unavailable. Registered NCI Joinpoint 6.1.0 output can be compared as optional validation but is not part of the primary estimator.")
    add_paragraph(doc,"Specification sensitivities allow one or three maximum breakpoints, require at least five years per segment, truncate at 2019, or fit on the rate scale. A UI-width-weighted fixed-knot sensitivity is also reported. The primary analysis remains the 1990-2023 log-rate model with zero to two breakpoints and at least four annual observations per segment.")
    add_heading(doc,"Descriptive trajectory contrasts")
    add_paragraph(doc,"For each comparison, we report each group's annualized endpoint change, the difference for group B minus group A, correlation between annual log changes, and whether endpoint directions match. These summaries do not test parallelism. Lag-1 residual autocorrelation and Durbin-Watson statistics diagnose whether the independent-error approximation is implausible.")
    add_heading(doc,"Shapley decomposition")
    add_paragraph(doc,"For all-age population N, age-share vector S, and age-specific rate vector R, reconstructed burden is B=N sum_a(S_a R_a). Each factor was replaced from baseline to endpoint in all 3! orders. A factor's contribution is the average marginal change across orders. The three contributions therefore sum exactly to B_end-B_start apart from floating-point tolerance; reconstructed endpoints are also checked against reported all-age numbers. The production age vector has 20 groups from 0-4 through 95+. A sensitivity collapses the finest available vector into 0-14, 15-49, 50-69, and 70+ and flags sign, magnitude-rank, or component shifts of at least 10% of total change.")
    add_heading(doc,"APC estimable functions")
    add_paragraph(doc,"The exact relation cohort=period-age prevents simultaneous identification of unrestricted linear age, period, and cohort effects. The model therefore contains an intercept, two identifiable linear trends, and nonlinear age, period, and cohort bases constrained to be orthogonal to intercept and linear trend. It does not estimate three unrestricted linear effects. Net drift, local drift, longitudinal age relative risks, period relative risks, and cohort relative risks are point estimates. Reference relative risks equal one. The primary window contains six five-year periods from 1994-1998 through 2019-2023; the sensitivity contains six periods from 1990-1994 through 2015-2019.")
    doc.add_page_break()
    add_heading(doc,"Uncertainty taxonomy")
    add_dataframe_table(doc,pd.DataFrame([
        {"Quantity":"Original GBD estimate","Interval":"Native 95% UI","Interpretation":"2.5th-97.5th percentiles of GBD draws"},
        {"Quantity":"Segment APC/AAPC","Interval":"None","Interpretation":"Descriptive point summary of annual GBD posterior means"},
        {"Quantity":"Ratios and changes","Interval":"None","Interpretation":"Point estimate; interval dominance is conservative only"},
        {"Quantity":"Decomposition component","Interval":"None","Interpretation":"Deterministic attribution using posterior means"},
    ]),[("Quantity","Quantity"),("Interval","Reported uncertainty"),("Interpretation","Interpretation")],font_size=8.5)
    out=out_dir/"statistical_methods_appendix.docx"; doc.save(out); return out


def build_gather(out_dir: Path, meta: dict):
    provisional=not bool(meta.get("submission_ready", False))
    doc=setup_document("GATHER reporting checklist","Guidelines for Accurate and Transparent Health Estimates Reporting",provisional)
    rows=[
        (1,"Define indicators, populations, and time periods","Methods: Outcomes; Study design","Complete"),
        (2,"List funding sources and funder roles","Declarations: Funding","Partial: author-confirmed funding and funder roles pending"),
        (3,"Describe source data access","Methods: Data source; Availability statement","Partial: verified retrieval metadata and repository link pending"),
        (4,"Identify all data sources","Supplementary Table S1","Partial until matching fine-age burden and population exports are supplied"),
        (5,"Describe data-source inclusion/exclusion","Methods: Analytical exclusions","Complete"),
        (6,"Report data-source characteristics","Supplementary Table S1","Partial: export IDs, queries, retrieval dates, and raw-file hashes incomplete"),
        (7,"Provide source data in machine-readable form","Provenance CSV and source archives","Partial subject to IHME terms and repository deposit"),
        (8,"Describe data processing","Methods; Statistical appendix","Complete"),
        (9,"Describe model and parameter selection","Methods: Trend/APC; Statistical appendix","Complete"),
        (10,"Describe covariates","No covariates used","Not applicable"),
        (11,"Describe uncertainty analysis","Methods: Comparisons; Statistical appendix","Partial: posterior draws unavailable; derived outputs are point estimates"),
        (12,"Describe analytical or statistical software","Methods: Reproducibility; README","Complete"),
        (13,"Provide source code","publication_study source files","Partial: persistent repository deposit and link pending"),
        (14,"Provide results in machine-readable form","CSV and XLSX outputs","Complete"),
        (15,"Report quantitative estimates with uncertainty","Results; Table 1; Figures 1 and 3","Partial: complete for native estimates only"),
        (16,"Make comparison groups explicit","Methods and all tables","Complete"),
        (17,"Discuss limitations of data and methods","Discussion: Strengths and limitations","Complete"),
        (18,"State interpretation and implications","Discussion and Conclusions","Complete"),
    ]
    add_dataframe_table(doc,pd.DataFrame(rows,columns=["Item","Requirement","Location","Status"]),[("Item","Item"),("Requirement","Requirement"),("Location","Manuscript location"),("Status","Status")],font_size=8.5)
    if provisional: add_paragraph(doc,"Checklist status does not override the requirement to supply matching fine-age GBD 2023 burden and official population exports, complete provenance sidecars, and a persistent repository link before submission.")
    out=out_dir/"GATHER_checklist.docx"; doc.save(out); return out


def load_tables(analysis: Path) -> dict[str,pd.DataFrame]:
    return {p.stem:pd.read_csv(p) for p in (analysis/"tables").glob("*.csv")}


def main():
    p=argparse.ArgumentParser(); p.add_argument("--analysis-dir",type=Path,required=True); p.add_argument("--output-dir",type=Path)
    args=p.parse_args(); analysis=args.analysis_dir; out=args.output_dir or analysis/"documents"; out.mkdir(parents=True,exist_ok=True)
    meta=json.loads((analysis/"build_metadata.json").read_text(encoding="utf-8")); tables=load_tables(analysis)
    paths=[build_manuscript(analysis,out,meta,tables),build_supplement(analysis,out,meta,tables),build_methods_appendix(out,meta),build_gather(out,meta)]
    (out/"document_manifest.json").write_text(json.dumps([str(x) for x in paths],indent=2),encoding="utf-8")
    print("\n".join(map(str,paths)))


if __name__=="__main__": main()
